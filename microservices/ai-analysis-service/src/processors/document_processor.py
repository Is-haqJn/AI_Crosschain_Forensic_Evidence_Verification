"""
Document Analysis Processor
Handles comprehensive document forensics and authenticity verification
"""

import asyncio
import hashlib
import time
from datetime import datetime
from pathlib import Path
from typing import Dict, Any, List, Optional
import PyPDF2
import docx
from docx import Document
from docx.oxml.ns import qn
import openpyxl
from loguru import logger

from ..schemas.analysis_schemas import (
    AnalysisRequest,
    DocumentAnalysisResult,
    DocumentAuthenticityResult,
    DocumentContentAnalysis
)
from ..models import get_model_manager
from ..config import get_settings

settings = get_settings()


class DocumentProcessor:
    """Advanced document analysis processor for forensic evidence"""
    
    def __init__(self):
        self.model_manager = get_model_manager()
        self.supported_formats = settings.DOCUMENT_ALLOWED_FORMATS
        self.max_size = settings.DOCUMENT_MAX_SIZE
        self.confidence_threshold = settings.DOCUMENT_CONFIDENCE_THRESHOLD
    
    async def analyze(self, file_path: str, request: AnalysisRequest) -> DocumentAnalysisResult:
        """
        Perform comprehensive document analysis
        
        Args:
            file_path: Path to the document file
            request: Analysis request details
            
        Returns:
            Complete document analysis results
        """
        start_time = time.time()
        
        try:
            # Run all analysis components in parallel with fault tolerance
            tasks = [
                self._analyze_authenticity(file_path),
                self._analyze_content(file_path),
                self._analyze_metadata(file_path),
                self._analyze_structure(file_path),
                self._check_plagiarism(file_path)
            ]

            results = await asyncio.gather(*tasks, return_exceptions=True)

            def _safe(idx: int, default):
                val = results[idx]
                return val if not isinstance(val, Exception) else default

            authenticity_result = _safe(0, DocumentAuthenticityResult(is_authentic=False, confidence=0.0, forgery_indicators=["analysis_error"], digital_signatures=[], creation_software=None))
            content_analysis = _safe(1, DocumentContentAnalysis(text_content="", language="unknown", word_count=0, character_count=0, readability_score=0.0, sensitive_information=[], classification="unknown"))
            metadata_analysis = _safe(2, {})
            structure_analysis = _safe(3, {"error": "structure_analysis_failed"})
            plagiarism_check = _safe(4, {"error": "plagiarism_check_failed"})

            processing_time = time.time() - start_time

            # Calculate overall confidence
            overall_confidence = self._calculate_overall_confidence(
                authenticity_result, content_analysis, metadata_analysis
            )

            return DocumentAnalysisResult(
                analysis_id=request.analysis_id,
                evidence_id=request.evidence_id,
                confidence_score=overall_confidence,
                processing_time=processing_time,
                model_version="1.0.0",
                authenticity_analysis=authenticity_result,
                content_analysis=content_analysis,
                metadata_analysis=metadata_analysis,
                structure_analysis=structure_analysis,
                plagiarism_check=plagiarism_check,
                metadata={
                    "file_path": file_path,
                    "analysis_timestamp": datetime.utcnow().isoformat(),
                    "processor_version": "1.0.0"
                }
            )

        except Exception as e:
            # Return neutral result on fatal error instead of raising
            processing_time = time.time() - start_time
            logger.error(f"Error analyzing document {file_path}: {e}")
            return DocumentAnalysisResult(
                analysis_id=request.analysis_id,
                evidence_id=request.evidence_id,
                confidence_score=0.0,
                processing_time=processing_time,
                model_version="1.0.0",
                authenticity_analysis=DocumentAuthenticityResult(is_authentic=False, confidence=0.0, forgery_indicators=["fatal_error"], digital_signatures=[], creation_software=None),
                content_analysis=DocumentContentAnalysis(text_content="", language="unknown", word_count=0, character_count=0, readability_score=0.0, sensitive_information=[], classification="unknown"),
                metadata_analysis={},
                structure_analysis={"error": "fatal_error"},
                plagiarism_check={"error": "fatal_error"},
                metadata={
                    "file_path": file_path,
                    "analysis_timestamp": datetime.utcnow().isoformat(),
                    "processor_version": "1.0.0",
                    "error": "document_analysis_failed"
                }
            )
    
    async def _analyze_authenticity(self, file_path: str) -> DocumentAuthenticityResult:
        """Analyze document authenticity using real AI models"""
        try:
            # Get document authenticity model
            model = await self.model_manager.get_model("document_authenticity_verifier")
            
            if model is None:
                logger.warning("Document authenticity model not loaded, using fallback")
                return await self._fallback_authenticity_analysis(file_path)
            
            # Use the real analysis method from the model
            if hasattr(model, 'analyze') and callable(model.get('analyze')):
                logger.info("Performing real AI-based document authenticity analysis")
                analysis_result = model['analyze'](file_path)
                
                # Convert the analysis result to our schema
                return DocumentAuthenticityResult(
                    is_authentic=analysis_result.get('is_authentic', True),
                    confidence=analysis_result.get('confidence_score', 0.8),
                    forgery_indicators=analysis_result.get('authenticity_indicators', []),
                    digital_signatures=[],  # Will be populated by other methods
                    creation_software=analysis_result.get('document_type', 'Unknown')
                )
            else:
                logger.warning("Model does not have analyze method, using fallback")
                return await self._fallback_authenticity_analysis(file_path)
            
        except Exception as e:
            logger.error(f"Authenticity analysis failed: {e}")
            return await self._fallback_authenticity_analysis(file_path)
    
    async def _fallback_authenticity_analysis(self, file_path: str) -> DocumentAuthenticityResult:
        """Fallback authenticity analysis using basic techniques"""
        try:
            # Basic file analysis
            file_info = await self._get_basic_file_info(file_path)
            
            # Simple heuristics for authenticity
            forgery_indicators = []
            
            # Check file modification time vs creation time
            if file_info.get("modified_after_creation", False):
                forgery_indicators.append("File modified after creation")
            
            # Check for suspicious metadata
            if file_info.get("suspicious_metadata", False):
                forgery_indicators.append("Suspicious metadata detected")
            
            # Basic authenticity assessment
            is_authentic = len(forgery_indicators) == 0
            confidence = 0.7 if is_authentic else 0.3
            
            return DocumentAuthenticityResult(
                is_authentic=is_authentic,
                confidence=confidence,
                forgery_indicators=forgery_indicators,
                digital_signatures=[],
                creation_software=file_info.get("creation_software")
            )
            
        except Exception as e:
            logger.error(f"Fallback authenticity analysis failed: {e}")
            return DocumentAuthenticityResult(
                is_authentic=False,
                confidence=0.0,
                forgery_indicators=["Analysis failed"],
                digital_signatures=[],
                creation_software=None
            )
    
    async def _analyze_content(self, file_path: str) -> DocumentContentAnalysis:
        """Analyze document content and extract text"""
        try:
            # Extract text content based on file type
            text_content = await self._extract_text_content(file_path)
            
            # Analyze content properties
            word_count = len(text_content.split())
            character_count = len(text_content)
            
            # Calculate readability score (simplified)
            readability_score = await self._calculate_readability_score(text_content)
            
            # Detect language
            language = await self._detect_language(text_content)
            
            # Detect sensitive information
            sensitive_information = await self._detect_sensitive_information(text_content)
            
            # Classify content
            classification = await self._classify_content(text_content)
            
            return DocumentContentAnalysis(
                text_content=text_content,
                language=language,
                word_count=word_count,
                character_count=character_count,
                readability_score=readability_score,
                sensitive_information=sensitive_information,
                classification=classification
            )
            
        except Exception as e:
            logger.error(f"Content analysis failed: {e}")
            return DocumentContentAnalysis(
                text_content="",
                language="unknown",
                word_count=0,
                character_count=0,
                readability_score=0.0,
                sensitive_information=[],
                classification="unknown"
            )
    
    async def _analyze_metadata(self, file_path: str) -> Dict[str, Any]:
        """Analyze document metadata"""
        try:
            metadata = {}
            file_path_obj = Path(file_path)
            file_ext = file_path_obj.suffix.lower()
            
            if file_ext == '.pdf':
                metadata = await self._extract_pdf_metadata(file_path)
            elif file_ext in ['.doc', '.docx']:
                metadata = await self._extract_word_metadata(file_path)
            elif file_ext in ['.xls', '.xlsx']:
                metadata = await self._extract_excel_metadata(file_path)
            else:
                metadata = await self._extract_basic_metadata(file_path)
            
            return metadata
            
        except Exception as e:
            logger.error(f"Metadata analysis failed: {e}")
            return {"error": str(e)}
    
    async def _analyze_structure(self, file_path: str) -> Dict[str, Any]:
        """Analyze document structure"""
        try:
            structure_info = {}
            file_path_obj = Path(file_path)
            file_ext = file_path_obj.suffix.lower()
            
            if file_ext == '.pdf':
                structure_info = await self._analyze_pdf_structure(file_path)
            elif file_ext in ['.doc', '.docx']:
                structure_info = await self._analyze_word_structure(file_path)
            elif file_ext in ['.xls', '.xlsx']:
                structure_info = await self._analyze_excel_structure(file_path)
            else:
                structure_info = {"type": "unknown", "structure": "basic"}
            
            return structure_info
            
        except Exception as e:
            logger.error(f"Structure analysis failed: {e}")
            return {"error": str(e)}
    
    async def _check_plagiarism(self, file_path: str) -> Dict[str, Any]:
        """Check document for plagiarism"""
        try:
            # Extract text content
            text_content = await self._extract_text_content(file_path)
            
            # Basic plagiarism detection (simplified)
            plagiarism_indicators = []
            similarity_score = 0.0
            
            # Check for common phrases and patterns
            common_phrases = [
                "Lorem ipsum dolor sit amet",
                "The quick brown fox jumps over the lazy dog",
                "To be or not to be, that is the question"
            ]
            
            for phrase in common_phrases:
                if phrase.lower() in text_content.lower():
                    plagiarism_indicators.append(f"Common phrase detected: {phrase}")
                    similarity_score += 0.1
            
            # Check for repeated content
            words = text_content.split()
            if len(words) > 0:
                unique_words = set(words)
                repetition_ratio = 1.0 - (len(unique_words) / len(words))
                if repetition_ratio > 0.3:
                    plagiarism_indicators.append("High content repetition detected")
                    similarity_score += repetition_ratio * 0.5
            
            return {
                "similarity_score": min(similarity_score, 1.0),
                "plagiarism_indicators": plagiarism_indicators,
                "is_plagiarized": similarity_score > 0.5,
                "analysis_method": "basic_pattern_detection"
            }
            
        except Exception as e:
            logger.error(f"Plagiarism check failed: {e}")
            return {"error": str(e)}
    
    # Helper methods for content extraction
    async def _extract_text_content(self, file_path: str) -> str:
        """Extract text content from document"""
        try:
            file_path_obj = Path(file_path)
            file_ext = file_path_obj.suffix.lower()
            
            if file_ext == '.pdf':
                return await self._extract_pdf_text(file_path)
            elif file_ext in ['.doc', '.docx']:
                return await self._extract_word_text(file_path)
            elif file_ext in ['.xls', '.xlsx']:
                return await self._extract_excel_text(file_path)
            elif file_ext == '.txt':
                return await self._extract_txt_text(file_path)
            else:
                return ""
                
        except Exception as e:
            logger.error(f"Text extraction failed: {e}")
            return ""
    
    async def _extract_pdf_text(self, file_path: str) -> str:
        """Extract text from PDF file"""
        try:
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
            return text.strip()
        except Exception as e:
            logger.error(f"PDF text extraction failed: {e}")
            return ""
    
    async def _extract_word_text(self, file_path: str) -> str:
        """Extract text from Word document"""
        try:
            doc = Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text.strip()
        except Exception as e:
            logger.error(f"Word text extraction failed: {e}")
            return ""
    
    async def _extract_excel_text(self, file_path: str) -> str:
        """Extract text from Excel file"""
        try:
            workbook = openpyxl.load_workbook(file_path)
            text = ""
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                for row in sheet.iter_rows():
                    row_text = " ".join([str(cell.value) for cell in row if cell.value])
                    if row_text:
                        text += row_text + "\n"
            return text.strip()
        except Exception as e:
            logger.error(f"Excel text extraction failed: {e}")
            return ""
    
    async def _extract_txt_text(self, file_path: str) -> str:
        """Extract text from plain text file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read().strip()
        except Exception as e:
            logger.error(f"TXT text extraction failed: {e}")
            return ""
    
    # Additional helper methods
    async def _extract_digital_signatures(self, file_path: str) -> List[Dict[str, Any]]:
        """Extract digital signatures from document"""
        return []
    
    async def _detect_creation_software(self, file_path: str) -> Optional[str]:
        """Detect software used to create the document"""
        try:
            file_ext = Path(file_path).suffix.lower()
            
            if file_ext == '.pdf':
                metadata = await self._extract_pdf_metadata(file_path)
                return metadata.get("creator") or metadata.get("producer")
            elif file_ext in ['.doc', '.docx']:
                metadata = await self._extract_word_metadata(file_path)
                return metadata.get("creator")
            else:
                return None
                
        except Exception as e:
            logger.error(f"Creation software detection failed: {e}")
            return None
    
    async def _detect_forgery_indicators(self, file_path: str) -> List[str]:
        """Detect indicators of document forgery"""
        return []
    
    async def _get_basic_file_info(self, file_path: str) -> Dict[str, Any]:
        """Get basic file information"""
        try:
            file_path_obj = Path(file_path)
            stat = file_path_obj.stat()
            
            return {
                "file_size": stat.st_size,
                "created": datetime.fromtimestamp(stat.st_ctime),
                "modified": datetime.fromtimestamp(stat.st_mtime),
                "modified_after_creation": stat.st_mtime > stat.st_ctime,
                "suspicious_metadata": False,
                "creation_software": None
            }
        except Exception as e:
            logger.error(f"Basic file info extraction failed: {e}")
            return {}
    
    # Analysis helper methods
    async def _calculate_readability_score(self, text: str) -> float:
        """Calculate readability score (simplified)"""
        try:
            if not text:
                return 0.0
            
            words = text.split()
            sentences = text.split('.')
            
            if len(words) == 0 or len(sentences) == 0:
                return 0.0
            
            avg_words_per_sentence = len(words) / len(sentences)
            avg_syllables_per_word = sum(len(word) for word in words) / len(words) / 3
            
            # Simplified Flesch Reading Ease formula
            score = 206.835 - (1.015 * avg_words_per_sentence) - (84.6 * avg_syllables_per_word)
            return max(0.0, min(100.0, score)) / 100.0
            
        except Exception as e:
            logger.error(f"Readability calculation failed: {e}")
            return 0.5
    
    async def _detect_language(self, text: str) -> str:
        """Detect document language (simplified)"""
        try:
            if not text:
                return "unknown"
            
            # Simple language detection based on common words
            english_words = ["the", "and", "or", "but", "in", "on", "at", "to", "for", "of", "with", "by"]
            spanish_words = ["el", "la", "de", "que", "y", "a", "en", "un", "es", "se", "no", "te", "lo", "le"]
            french_words = ["le", "la", "de", "et", "à", "un", "il", "que", "ne", "se", "ce", "pas", "tout"]
            
            text_lower = text.lower()
            
            english_count = sum(1 for word in english_words if word in text_lower)
            spanish_count = sum(1 for word in spanish_words if word in text_lower)
            french_count = sum(1 for word in french_words if word in text_lower)
            
            if english_count > spanish_count and english_count > french_count:
                return "english"
            elif spanish_count > french_count:
                return "spanish"
            elif french_count > 0:
                return "french"
            else:
                return "unknown"
                
        except Exception as e:
            logger.error(f"Language detection failed: {e}")
            return "unknown"
    
    async def _detect_sensitive_information(self, text: str) -> List[str]:
        """Detect sensitive information in text"""
        try:
            sensitive_info = []
            
            # Simple pattern matching for sensitive information
            import re
            patterns = {
                "email": r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b',
                "phone": r'\b\d{3}-\d{3}-\d{4}\b',
                "ssn": r'\b\d{3}-\d{2}-\d{4}\b'
            }
            
            for info_type, pattern in patterns.items():
                matches = re.findall(pattern, text)
                if matches:
                    sensitive_info.append(f"{info_type}: {len(matches)} instances found")
            
            return sensitive_info
            
        except Exception as e:
            logger.error(f"Sensitive information detection failed: {e}")
            return []
    
    async def _classify_content(self, text: str) -> str:
        """Classify document content type"""
        try:
            if not text:
                return "empty"
            
            text_lower = text.lower()
            
            # Simple content classification
            if any(word in text_lower for word in ["contract", "agreement", "terms", "conditions"]):
                return "legal"
            elif any(word in text_lower for word in ["invoice", "bill", "payment", "receipt"]):
                return "financial"
            elif any(word in text_lower for word in ["report", "analysis", "data", "statistics"]):
                return "report"
            elif any(word in text_lower for word in ["email", "message", "correspondence"]):
                return "correspondence"
            else:
                return "general"
                
        except Exception as e:
            logger.error(f"Content classification failed: {e}")
            return "unknown"
    
    # Metadata extraction methods
    async def _extract_pdf_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract metadata from PDF file"""
        try:
            metadata = {}
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                if pdf_reader.metadata:
                    metadata = {
                        "title": pdf_reader.metadata.get("/Title"),
                        "author": pdf_reader.metadata.get("/Author"),
                        "subject": pdf_reader.metadata.get("/Subject"),
                        "creator": pdf_reader.metadata.get("/Creator"),
                        "producer": pdf_reader.metadata.get("/Producer"),
                        "creation_date": pdf_reader.metadata.get("/CreationDate"),
                        "modification_date": pdf_reader.metadata.get("/ModDate"),
                        "page_count": len(pdf_reader.pages)
                    }
            return metadata
        except Exception as e:
            logger.error(f"PDF metadata extraction failed: {e}")
            return {}
    
    async def _extract_word_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract metadata from Word document"""
        try:
            doc = Document(file_path)
            core_props = doc.core_properties
            
            metadata = {
                "title": core_props.title,
                "author": core_props.author,
                "subject": core_props.subject,
                "creator": core_props.creator,
                "created": core_props.created.isoformat() if core_props.created else None,
                "modified": core_props.modified.isoformat() if core_props.modified else None,
                "paragraph_count": len(doc.paragraphs)
            }
            return metadata
        except Exception as e:
            logger.error(f"Word metadata extraction failed: {e}")
            return {}
    
    async def _extract_excel_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract metadata from Excel file"""
        try:
            workbook = openpyxl.load_workbook(file_path)
            props = workbook.properties
            
            metadata = {
                "title": props.title,
                "author": props.creator,
                "created": props.created.isoformat() if props.created else None,
                "modified": props.modified.isoformat() if props.modified else None,
                "sheet_count": len(workbook.sheetnames),
                "sheet_names": workbook.sheetnames
            }
            return metadata
        except Exception as e:
            logger.error(f"Excel metadata extraction failed: {e}")
            return {}
    
    async def _extract_basic_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract basic file metadata"""
        try:
            file_path_obj = Path(file_path)
            stat = file_path_obj.stat()
            
            return {
                "file_name": file_path_obj.name,
                "file_size": stat.st_size,
                "created": datetime.fromtimestamp(stat.st_ctime).isoformat(),
                "modified": datetime.fromtimestamp(stat.st_mtime).isoformat(),
                "extension": file_path_obj.suffix
            }
        except Exception as e:
            logger.error(f"Basic metadata extraction failed: {e}")
            return {}
    
    # Structure analysis methods
    async def _analyze_pdf_structure(self, file_path: str) -> Dict[str, Any]:
        """Analyze PDF document structure"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)

                # Section count: use PDF outlines/bookmarks when available
                def _flatten_outline(items) -> int:
                    try:
                        count = 0
                        for item in items:
                            if isinstance(item, list):
                                count += _flatten_outline(item)
                            else:
                                count += 1
                        return count
                    except Exception:
                        return 0

                try:
                    outlines = getattr(pdf_reader, 'outline', None)
                    if outlines is None and hasattr(pdf_reader, 'get_outlines'):
                        outlines = list(pdf_reader.get_outlines())
                    section_count = _flatten_outline(outlines) if outlines else 0
                except Exception:
                    section_count = 0

                # Image count: count XObjects of subtype Image (including nested Forms)
                def _count_images_from_xobject(xobj) -> int:
                    image_total = 0
                    try:
                        for obj in xobj.values():
                            try:
                                resolved = obj.get_object()
                            except Exception:
                                resolved = obj
                            subtype = resolved.get('/Subtype')
                            if subtype == '/Image':
                                image_total += 1
                            elif subtype == '/Form':
                                resources = resolved.get('/Resources')
                                if resources and '/XObject' in resources:
                                    image_total += _count_images_from_xobject(resources['/XObject'])
                    except Exception:
                        pass
                    return image_total

                image_count = 0
                try:
                    for page in pdf_reader.pages:
                        resources = page.get('/Resources')
                        if resources and '/XObject' in resources:
                            image_count += _count_images_from_xobject(resources['/XObject'])
                except Exception:
                    image_count = 0

                return {
                    "type": "pdf",
                    "page_count": len(pdf_reader.pages),
                    "section_count": section_count,
                    "table_count": 0,  # Table detection in PDFs requires specialized libs; default to 0
                    "image_count": image_count,
                    "has_bookmarks": section_count > 0,
                    "is_encrypted": pdf_reader.is_encrypted,
                    "structure": "standard_pdf"
                }
        except Exception as e:
            logger.error(f"PDF structure analysis failed: {e}")
            return {"error": str(e)}
    
    async def _analyze_word_structure(self, file_path: str) -> Dict[str, Any]:
        """Analyze Word document structure"""
        try:
            doc = Document(file_path)

            # Section count from document sections
            section_count = len(doc.sections)

            # Estimate page count: count explicit page breaks (best-effort) + 1
            try:
                page_breaks = 0
                NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
                for paragraph in doc.paragraphs:
                    for run in paragraph.runs:
                        el = run._element
                        for br in el.xpath('.//w:br', namespaces=NS):
                            br_type = br.get(qn('w:type'))
                            if br_type == 'page':
                                page_breaks += 1
                page_count = max(1, page_breaks + 1)
            except Exception:
                page_count = 1

            # Count images via document relationships (captures inline and floating)
            try:
                image_rels = [r for r in doc.part.rels.values() if 'image' in r.reltype]
                image_count = len(image_rels)
            except Exception:
                image_count = len(getattr(doc, 'inline_shapes', []))

            table_count = len(doc.tables)

            return {
                "type": "word",
                "page_count": page_count,
                "section_count": section_count,
                "table_count": table_count,
                "image_count": image_count,
                "paragraph_count": len(doc.paragraphs),
                "structure": "standard_word"
            }
        except Exception as e:
            logger.error(f"Word structure analysis failed: {e}")
            return {"error": str(e)}
    
    async def _analyze_excel_structure(self, file_path: str) -> Dict[str, Any]:
        """Analyze Excel document structure"""
        try:
            workbook = openpyxl.load_workbook(file_path)
            
            total_cells = 0
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                total_cells += sheet.max_row * sheet.max_column
            # Tables per sheet
            try:
                table_count = sum(len(getattr(workbook[s], 'tables', {})) for s in workbook.sheetnames)
            except Exception:
                table_count = 0
            # Images per sheet (best-effort; openpyxl stores on private attr)
            try:
                image_count = 0
                for sheet_name in workbook.sheetnames:
                    sheet = workbook[sheet_name]
                    image_count += len(getattr(sheet, '_images', []))
            except Exception:
                image_count = 0
            
            return {
                "type": "excel",
                "sheet_count": len(workbook.sheetnames),
                "sheet_names": workbook.sheetnames,
                "total_cells": total_cells,
                "page_count": 0,
                "section_count": len(workbook.sheetnames),
                "table_count": table_count,
                "image_count": image_count,
                "structure": "standard_excel"
            }
        except Exception as e:
            logger.error(f"Excel structure analysis failed: {e}")
            return {"error": str(e)}
    
    def _calculate_overall_confidence(
        self,
        authenticity_result: DocumentAuthenticityResult,
        content_analysis: DocumentContentAnalysis,
        metadata_analysis: Dict[str, Any]
    ) -> float:
        """Calculate overall analysis confidence"""
        # Weighted average of different analysis components
        weights = {
            "authenticity": 0.4,
            "content": 0.3,
            "metadata": 0.3
        }
        
        confidence = (
            authenticity_result.confidence * weights["authenticity"] +
            (1.0 if content_analysis.word_count > 0 else 0.0) * weights["content"] +
            (1.0 if metadata_analysis else 0.0) * weights["metadata"]
        )
        
        return min(confidence, 1.0)